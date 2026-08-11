"""Informe en Word dels desplacaments dels equips del Club Billar Banyoles."""

import json
import sys
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

SP = Path(sys.argv[1])
d = json.loads((SP / "rows.json").read_text(encoding="utf-8"))
ROWS = [r for r in d["rows"] if r["estat"] == "verificat"]
TOTS = [r for r in d["rows"] if r["estat"] in ("verificat", "incompareixenca")]
PROJ, ADRECES, P = d["projeccio"], d["adreces"], d["params"]
INCOMP, CASA = d["incompar"], d["a_casa"]
# Amb l'argument "fiscal" el document s'agrupa per any natural en comptes de per temporada.
FISCAL = len(sys.argv) > 2 and sys.argv[2] == "fiscal"


def periode(r):
    return r["data"][:4] if FISCAL else r["temporada"]


ETIQ = "Any" if FISCAL else "Temporada"
SEASONS = sorted({periode(r) for r in ROWS})

VERD = RGBColor(0x1F, 0x5F, 0x45)
GRIS = RGBColor(0x55, 0x60, 0x5A)
DIETA = f'{P["dieta"]:.2f}'.replace(".", ",")
LLINDAR = P["llindar_km"]
CONSUM = str(P["consum"]).replace(".", ",")


def km(v):
    return f"{v:,.1f}".replace(",", "@").replace(".", ",").replace("@", ".")


def eur(v, dec=2):
    return f"{v:,.{dec}f}".replace(",", "@").replace(".", ",").replace("@", ".") + " €"


def dt(iso):
    y, m, day = iso.split("-")
    return f"{day}/{m}/{y[2:]}"


def comp(r):
    return "Copa Catalana" if r["tipus"] == "copa" else "Lliga Tres Bandes"


def ordre(r):
    return ("Z" if r["tipus"] == "copa" else r["equip"], r["data"])


DIETA_4 = eur(4 * P["dieta"]).replace(" €", "")
DIETA_3 = eur(3 * P["dieta"]).replace(" €", "")


doc = Document()
st = doc.styles["Normal"]
st.font.name = "Calibri"
st.font.size = Pt(10)
for _sec in doc.sections:
    _sec.top_margin = _sec.bottom_margin = Cm(2)
    _sec.left_margin = _sec.right_margin = Cm(2)


def ombreja(cell, hexcolor):
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(sh)


def titol(text, mida=15, space_before=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(mida)
    r.font.color.rgb = VERD
    return p


def text(t, mida=10, cursiva=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(t)
    r.font.size = Pt(mida)
    r.italic = cursiva
    return p


def nova_seccio(horitzontal: bool):
    """Comenca pagina nova amb l'orientacio demanada i retorna l'amplada util en cm."""
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    a4_curt, a4_llarg = Cm(21), Cm(29.7)
    if horitzontal:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = a4_llarg, a4_curt
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width, sec.page_height = a4_curt, a4_llarg
    sec.top_margin = sec.bottom_margin = Cm(1.6)
    sec.left_margin = sec.right_margin = Cm(1.6)
    return (sec.page_width - sec.left_margin - sec.right_margin) / 360000


def _no_partir(row):
    """Evita que una fila es parteixi entre dues pagines."""
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:cantSplit")
    trPr.append(e)


def _repeteix_capcalera(row):
    """Repeteix la fila de capcalera a dalt de cada pagina."""
    trPr = row._tr.get_or_add_trPr()
    e = OxmlElement("w:tblHeader")
    e.set(qn("w:val"), "true")
    trPr.append(e)


def _posa(cell, valor, amplada, dreta, mida, bold=False, fons=None):
    cell.text = ""
    cell.width = Cm(amplada)
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    if dreta:
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(str(valor))
    r.font.size = Pt(mida)
    r.bold = bold
    if fons:
        ombreja(cell, fons)
    return r


AMPLE_UTIL = [17.0]


def taula(caps, files, amples, numeriques=(), totals=None, mida=8):
    # Escala les amplades perque la taula ompli exactament la caixa de text.
    factor = AMPLE_UTIL[0] / sum(amples)
    amples = [a * factor for a in amples]
    t = doc.add_table(rows=1, cols=len(caps))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    for i, a in enumerate(amples):
        t.columns[i].width = Cm(a)
    _repeteix_capcalera(t.rows[0])
    _no_partir(t.rows[0])
    for i, c in enumerate(caps):
        r = _posa(t.rows[0].cells[i], c, amples[i], i in numeriques, mida,
                  bold=True, fons="1F5F45")
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for f in files:
        row = t.add_row()
        _no_partir(row)
        for i, v in enumerate(f):
            _posa(row.cells[i], v, amples[i], i in numeriques, mida)
    if totals:
        row = t.add_row()
        _no_partir(row)
        for i, v in enumerate(totals):
            _posa(row.cells[i], v, amples[i], i in numeriques, mida,
                  bold=True, fons="E6EFE9")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return t


# ------------------------------------------------------------------ portada
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
p.add_run().add_picture(str(SP / "escut-cbb-900.png"), height=Cm(3.2))

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CLUB BILLAR BANYOLES")
r.bold = True
r.font.size = Pt(13)
r.font.color.rgb = VERD
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Carrer de l'Abeurador, 10 · 17820 Banyoles")
r.font.size = Pt(9)
r.font.color.rgb = GRIS

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(16)
r = p.add_run("DESPLAÇAMENTS DELS EQUIPS EN COMPETICIÓ OFICIAL")
r.bold = True
r.font.size = Pt(17)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Anys 2014 a 2026" if FISCAL else "Temporades 2014-2015 a 2025-2026")
r.font.size = Pt(12)
r.font.color.rgb = GRIS
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(22)
r = p.add_run("Memòria justificativa per a sol·licitud d'ajut econòmic")
r.font.size = Pt(10)
r.italic = True
r.font.color.rgb = GRIS

taula(
    ["Concepte", "Valor"],
    [["Desplaçaments a competició oficial", f"{len(ROWS)}"],
     ["Quilòmetres recorreguts (anada i tornada)", km(sum(r["km_total"] for r in ROWS))],
     ["Cost de quilometratge", eur(sum(r["import_km_eur"] for r in ROWS))],
     ["Cost de dietes de manutenció", eur(sum(r["import_dietes_eur"] for r in ROWS))],
     ["Clubs visitats", f"{len({r['seu_club'] for r in ROWS})}"]],
    amples=[10.6, 6.0], numeriques=(1,),
    totals=["COST TOTAL DEL PERÍODE", eur(sum(r["import_total_eur"] for r in ROWS))])

# ------------------------------------------------------------------ objecte
titol("1. Objecte i abast")
text("Aquest document detalla els desplaçaments realitzats pels equips del Club Billar "
     "Banyoles per disputar competicions oficials per equips de la Federació Catalana de "
     "Billar" + (" entre l'1 de gener de 2014 i el 31 de desembre de 2026, agrupats per any "
     "natural, i en quantifica el cost. Cal tenir present que la temporada esportiva va de "
     "setembre a maig, de manera que cada any natural en recull el tram final d'una i el "
     "començament de la següent." if FISCAL else
     " durant les temporades 2014-2015 a 2025-2026, i en quantifica el cost."))
text("S'hi inclouen la Lliga Catalana de Tres Bandes —fase regular, fases finals i "
     "promocions d'ascens— i la Copa Catalana per equips. No s'hi inclouen les proves "
     "individuals. Només s'hi comptabilitzen els desplaçaments per jugar fora de casa, "
     "sempre d'anada i tornada des de la seu del club.")
text("Cada desplaçament es detalla amb data, equip, competició, divisió, grup i club "
     "visitat, de manera que es pot contrastar un a un amb el portal públic de la "
     "Federació Catalana de Billar.")

titol("2. Criteris de valoració econòmica")
text(f"El quilometratge es valora amb el barem d'indemnització per ús de vehicle "
     f"particular exempt de gravamen a l'impost sobre la renda de les persones físiques, "
     f"aplicat amb la tarifa vigent el dia de cada partit: 0,19 €/km fins al 16 de juliol "
     f"de 2023 i 0,26 €/km a partir del 17 de juliol de 2023, segons l'Ordre HFP/792/2023, "
     f"de 12 de juliol. No s'hi han afegit peatges ni despeses d'aparcament.")
text(f"A cada desplaçament de més de {LLINDAR} quilòmetres d'anada s'hi suma una dieta de "
     f"manutenció de migdia de {DIETA} € per jugador, import que aplica el club i que es "
     f"manté per sota dels 26,67 € que l'article 9 del Reglament de l'impost fixa com a "
     f"límit exempt de gravamen. Els equips es desplacen amb quatre jugadors a la lliga i "
     f"amb tres a la Copa, on cada equip presenta tres jugadors: són {DIETA_4} € de dietes "
     f"per desplaçament de lliga i {DIETA_3} € per jornada de Copa.")

# ------------------------------------------------------------------ resum
# Sense salt forcat: amb un titol mes llarg o un text mes extens hi quedava mig full buit.
p3 = titol(f"3. Resum per {'any fiscal' if FISCAL else 'temporada'}")
p3.paragraph_format.keep_with_next = True
# La 2020-2021 no te desplacaments pero es llista igualment, al seu lloc.
files = []
for s in SEASONS if FISCAL else sorted(set(SEASONS) | {"2020-2021"}):
    c = [r for r in ROWS if periode(r) == s]
    if not c:
        files.append([s, "cap", "0", "0,0", eur(0), eur(0), eur(0)])
        continue
    eq = sorted({r["equip"].replace("Banyoles ", "").replace('"', "")
                 for r in c if r["tipus"] == "regular"})
    files.append([s, ", ".join(eq), str(len(c)), km(sum(r["km_total"] for r in c)),
                  eur(sum(r["import_km_eur"] for r in c)),
                  eur(sum(r["import_dietes_eur"] for r in c)),
                  eur(sum(r["import_total_eur"] for r in c))])
taula([ETIQ, "Equips", "Despl.", "km", "Quilometratge", "Dietes", "Total"],
      files, amples=[2.3, 2.0, 1.5, 2.0, 2.6, 2.3, 2.4], numeriques=(2, 3, 4, 5, 6),
      totals=["TOTAL", "", str(len(ROWS)), km(sum(r["km_total"] for r in ROWS)),
              eur(sum(r["import_km_eur"] for r in ROWS)),
              eur(sum(r["import_dietes_eur"] for r in ROWS)),
              eur(sum(r["import_total_eur"] for r in ROWS))])
if not FISCAL:
    text("De les dotze temporades del període, la 2020-2021 és l'única sense cap "
         "desplaçament: el club no hi va inscriure cap equip. Va ser la temporada "
         "afectada per la COVID-19, en què la competició no va començar fins al 13 de "
         "març de 2021 i tampoc no es va disputar la Copa Catalana per equips.",
         cursiva=True)
else:
    text("El club no va tenir equips en competició la temporada 2020-2021, afectada per la "
         "COVID-19, i per això entre el març de 2020 i el setembre de 2021 no hi consta cap "
         "desplaçament.", cursiva=True)

# ------------------------------------------------------------------ detall
AMPLE_UTIL[0] = nova_seccio(horitzontal=True)
titol(f"4. Detall per {'any fiscal' if FISCAL else 'temporada'}", space_before=0)
for s in SEASONS:
    sel = sorted([r for r in ROWS if periode(r) == s], key=ordre)
    ptit = titol(f"4.{SEASONS.index(s) + 1}. {ETIQ} {s}", mida=12, space_before=0)
    if SEASONS.index(s) > 0:
        # Salt de pagina al titol, no com a paragraf a part: aixi no queda cap full en blanc
        # quan la taula anterior ja arribava al final de la pagina.
        ptit.paragraph_format.page_break_before = True
    agg = defaultdict(lambda: [0, 0.0, 0.0, 0.0])
    for r in sel:
        k = (ordre(r)[0], r["equip"], comp(r))
        agg[k][0] += 1
        agg[k][1] += r["km_total"]
        agg[k][2] += r["import_km_eur"]
        agg[k][3] += r["import_dietes_eur"]
    taula(["Equip", "Competició", "Despl.", "km", "Quilometratge", "Dietes", "Total"],
          [[k[1], k[2], str(n), km(t), eur(q), eur(di), eur(q + di)]
           for k, (n, t, q, di) in sorted(agg.items())],
          amples=[2.3, 3.2, 1.5, 2.0, 2.5, 2.2, 2.4], numeriques=(2, 3, 4, 5, 6),
          totals=["TOTAL", "", str(len(sel)), km(sum(r["km_total"] for r in sel)),
                  eur(sum(r["import_km_eur"] for r in sel)),
                  eur(sum(r["import_dietes_eur"] for r in sel)),
                  eur(sum(r["import_total_eur"] for r in sel))])
    files, anterior, hi_ha_asterisc = [], None, False
    for r in sorted([x for x in TOTS if periode(x) == s], key=ordre):
        bloc = ordre(r)[0]
        etiqueta = r["equip"] if bloc != anterior else ""
        anterior = bloc
        nj = r["estat"] == "incompareixenca"
        hi_ha_asterisc = hi_ha_asterisc or nj
        files.append([etiqueta, dt(r["data"]) + (" *" if nj else ""), comp(r), r["divisio"],
                      r["grup"], r["seu_club"],
                      km(r["km_total"]), eur(r["tarifa_eur_km"]), eur(r["import_km_eur"]),
                      str(r["jugadors"]),
                      DIETA + " €" if r["te_dieta"] else "—",
                      eur(r["import_dietes_eur"]), eur(r["import_total_eur"])])
    taula(["Equip", "Data", "Competició", "Divisió", "Grup", "Club on es juga",
           "km", "€/km", "Quilom.", "Jug.", "€/dieta", "Dietes", "Total"],
          files, amples=[1.9, 1.3, 1.8, 1.5, 1.3, 2.6, 1.2, 1.0, 1.4, 0.8, 1.1, 1.3, 1.4],
          numeriques=(6, 7, 8, 9, 10, 11, 12),
          totals=["TOTAL", "", "", "", "", "", km(sum(r["km_total"] for r in sel)),
                  "", eur(sum(r["import_km_eur"] for r in sel)), "", "",
                  eur(sum(r["import_dietes_eur"] for r in sel)),
                  eur(sum(r["import_total_eur"] for r in sel))])
    if hi_ha_asterisc:
        text("* Encontre no disputat per incompareixença: consta al calendari federatiu "
             "però no es va jugar, de manera que no genera desplaçament ni computa.",
             mida=8, cursiva=True)

# ------------------------------------------------------------------ projeccio
titol("5. Despesa estimada per a la temporada 2026-2027", space_before=0)
text("Els tres equips del club ja tenen divisió, però no grup assignats per a la temporada "
     "vinent. La composició dels grups és estimada a partir de les classificacions de l'any "
     "passat i pot tenir modificacions. Com que cada grup es juga a doble volta, el nombre "
     "de desplaçaments i les seus són coneguts d'entrada, encara que el calendari no "
     "estigui publicat: un viatge a cada un dels altres clubs del grup.")
text("Aquest apartat és una previsió i no forma part dels imports justificats de les "
     "temporades anteriors, de manera que no s'hi ha de sumar. No inclou la Copa Catalana, "
     "perquè encara no se sap si el club hi participarà. El quilometratge s'hi valora tot "
     "a 0,26 €/km, el barem vigent.", cursiva=True)
blocs_p = {}
for r in PROJ:
    blocs_p.setdefault(r["equip"], []).append(r)
taula(["Equip", "Competició", "Despl.", "km", "Quilometratge", "Dietes", "Total"],
      [[e, "Lliga Tres Bandes", str(len(b)), km(sum(x["km_total"] for x in b)),
        eur(sum(x["import_km_eur"] for x in b)),
        eur(sum(x["import_dietes_eur"] for x in b)),
        eur(sum(x["import_total_eur"] for x in b))] for e, b in sorted(blocs_p.items())],
      amples=[2.3, 3.2, 1.5, 2.0, 2.5, 2.2, 2.4], numeriques=(2, 3, 4, 5, 6),
      totals=["TOTAL", "", str(len(PROJ)), km(sum(r["km_total"] for r in PROJ)),
              eur(sum(r["import_km_eur"] for r in PROJ)),
              eur(sum(r["import_dietes_eur"] for r in PROJ)),
              eur(sum(r["import_total_eur"] for r in PROJ))])

files, anterior = [], None
for e, b in sorted(blocs_p.items()):
    for r in b:
        primera = e != anterior
        anterior = e
        files.append([e if primera else "", r["divisio"] if primera else "",
                      r["grup"] if primera else "",
                      f'{r["rival_club"]} {r["rival_equip"]}'.strip(), r["municipi"],
                      km(r["km_total"]), eur(r["tarifa_eur_km"]), eur(r["import_km_eur"]),
                      str(r["jugadors"]) if r["te_dieta"] else "—",
                      DIETA + " €" if r["te_dieta"] else "—",
                      eur(r["import_dietes_eur"]) if r["te_dieta"] else "—",
                      eur(r["import_total_eur"])])
taula(["Equip", "Divisió", "Grup", "Rival", "Municipi", "km", "€/km", "Quilom.", "Jug.",
       "€/dieta", "Dietes", "Total"],
      files, amples=[1.9, 1.7, 1.0, 2.9, 2.5, 1.2, 1.0, 1.4, 0.8, 1.1, 1.3, 1.4],
      numeriques=(5, 6, 7, 8, 9, 10, 11),
      totals=["TOTAL", "", "", f"{len(PROJ)} desplaçaments", "",
              km(sum(r["km_total"] for r in PROJ)), "",
              eur(sum(r["import_km_eur"] for r in PROJ)), "", "",
              eur(sum(r["import_dietes_eur"] for r in PROJ)),
              eur(sum(r["import_total_eur"] for r in PROJ))])

# ------------------------------------------------------------------ distancies
titol("6. Distàncies de referència")
text("Una sola distància per club, aplicada a tots els desplaçaments que s'hi han fet. "
     "L'adreça és la que consta al directori oficial de clubs de la Federació Catalana de "
     "Billar. L'origen de tots els trajectes és el Club Billar Banyoles, carrer de "
     "l'Abeurador, 10, de Banyoles.")
dist = {}
for r in ROWS + PROJ:
    dist.setdefault(r.get("seu_club") or r["rival_club"], r)
files = []
for nom, r in sorted(dist.items(), key=lambda x: x[1]["km_anada"]):
    n = sum(1 for x in ROWS if x["seu_club"] == nom)
    files.append([nom, ADRECES.get(nom, ""), r["municipi"], km(r["km_anada"]),
                  km(r["km_total"]), str(n) if n else "—"])
taula(["Club", "Adreça", "Municipi", "km anada", "Anada i tornada", "Visites"],
      files, amples=[3.2, 5.4, 2.9, 1.7, 2.1, 1.3], numeriques=(3, 4, 5))
text("La columna de visites recull els desplaçaments efectivament realitzats en el període "
     "justificat. Els clubs amb un guió només apareixen a la projecció de la temporada "
     "2026-2027.", cursiva=True)

# ------------------------------------------------------------------ annexos
AMPLE_UTIL[0] = nova_seccio(horitzontal=False)
titol("7. Metodologia i fonts", space_before=0)
metode = [
    ("Calendari i resultats", "Portal públic de la Federació Catalana de Billar "
     "(www.fcbillar.cat). S'han recorregut, per a cada temporada, totes les divisions i "
     "tots els grups on hi ha hagut un equip del club, jornada per jornada, i s'ha "
     "comprovat que el calendari recollit coincideix amb el publicat. De la 2021-2022 "
     "ençà s'ha contrastat la base de dades del club amb el portal, grup a grup i "
     "jornada a jornada (20 grups, 179 encontres programats i cap divergència de "
     "data); de la 2014-2015 a la 2019-2020, en què la base de dades del club no és "
     "completa, les dades s'han pres directament del portal, on hi consten 16 grups i "
     "148 encontres amb data."),
    ("Direcció dels desplaçaments", "Es compta desplaçament quan l'equip del Banyoles "
     "figura com a visitant."),
    ("Seu de joc", "A la lliga regular, el local de joc és el del club local de l'encontre. "
     "A la Copa, on la seu és la del club responsable de cada grup."),
    ("Adreces", "Directori oficial de clubs de la Federació Catalana de Billar, unificades "
     "a un sol criteri tipogràfic sense alterar-ne el contingut."),
    ("Coordenades", "Geocodificades amb Nominatim sobre dades d'OpenStreetMap i revisades "
     "una a una. La major part dels clubs estan situats a nivell de carrer o de portal; "
     "Canet de Mar, Sant Feliu de Codines i Cardona, a nivell de municipi, perquè el "
     "carrer indicat no consta a la cartografia, cosa que en trajectes de més de 70 km "
     "suposa una diferència inferior a l'1 %."),
    ("Distàncies", "Calculades amb OSRM, motor d'encaminament obert sobre la xarxa viària "
     "d'OpenStreetMap, perfil de vehicle i ruta més ràpida. Cada tram és la distància real "
     "per carretera entre les dues adreces, multiplicada per dos. El càlcul és reproduïble: "
     "amb les mateixes coordenades, qualsevol pot repetir-lo i obtenir la mateixa xifra."),
    ("Quilometratge", "Barem exempt de gravamen a l'IRPF vigent el dia de cada partit "
     "(Ordre HFP/792/2023 des del 17 de juliol de 2023)."),
    ("Dietes", "Article 9 del Reglament de l'IRPF, manutenció sense pernoctació en "
     f"territori espanyol. S'apliquen als desplaçaments de més de {LLINDAR} km d'anada."),
]
for c, t in metode:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f"{c}. ")
    r.bold = True
    r.font.size = Pt(10)
    r = p.add_run(t)
    r.font.size = Pt(10)

out = SP / ("Informe_desplacaments_CB_Banyoles_any_fiscal.docx" if FISCAL
            else "Informe_desplacaments_CB_Banyoles.docx")
doc.save(out)
print(f"{out.name}: {len(ROWS)} desplaçaments, {len(PROJ)} de projecció, "
      f"{len(doc.tables)} taules")
